from re import L
from docx import Document
from pathlib import Path
from errors import DocumentNotOpenError
import os


class WordExecutor:
    def __init__(self):
        self._command_map = {
            "new_document" : self._word_new_document,
            "open_document" : self._open_document,
            "new_paragraph" : self._word_new_paragraph,
            "show_document" : self._show_document
        }

    def execute(self, command: str, *args):
        if command not in self._command_map:
            raise Exception(f"Команда {command} не найдена")
        return self._command_map[command](*args)


    def _word_new_document(self, file_name: str):
        path = (Path.home() / "Desktop" / f"{file_name}.docx")
        self.doc_name = str(path)
        self.doc = Document()
        self.doc.save(self.doc_name)

    def _open_document(self, file_name: str):
        path = (Path.home() / "Desktop" / f"{file_name}.docx")
        self.doc_name = str(path)
        self.doc = Document(str(path))

    def _show_document(self, file_name: str):
        path = (Path.home() / "Desktop" / f"{file_name}.docx")
        if not os.path.exists(str(path)):
            raise DocumentNotOpenError()
        os.system(f"start {str(path)}")

    def _word_new_paragraph(self, text: str):
        if not self.doc:
            raise DocumentNotOpenError()
        if "курсив" in text:
            self.doc.add_paragraph(text).italic = True #type: ignore
        elif "жирный" in text:
            self.doc.add_paragraph(text).bold = True #type: ignore
        elif "подчеркнутый" in text:
            self.doc.add_paragraph(text).underline = True #type: ignore
        elif "зачеркнутый" in text:
            self.doc.add_paragraph(text).strike = True #type: ignore
        else:
            self.doc.add_paragraph(text)
        self.doc.save(self.doc_name)

    def _new_head(self, text: str):
        if not self.doc:
            raise DocumentNotOpenError()
        self.doc.add_heading(text)
        self.doc.save(self.doc_name)

    def _set_font(self, text: str):
        if not self.doc:
            raise DocumentNotOpenError()
        self.style = self.doc.styles['Normal']
        font = style.font.name = text # type: ignore

    def _set_size(self, size: int):
        if not self.doc:
            raise DocumentNotOpenError()
        style = self.doc.styles['Normal']
        font_size = style.font.size = size # type: ignore